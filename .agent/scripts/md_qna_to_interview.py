#!/usr/bin/env python3
"""
md_qna_to_interview.py
Convert file questions_tracking_*.md thành DOCX phỏng vấn chuyên nghiệp
theo template chuẩn của BA Team (màu 1F4E79, 3-column question table).

Author: M2MBA
Version: 1.0.0
Usage: py md_qna_to_interview.py <input.md> [output.docx]
"""

import sys
import os
import re
import subprocess

def install_package(package):
    try:
        __import__(package.replace('-', '_'))
    except ImportError:
        print(f"Installing '{package}'...")
        subprocess.check_call([sys.executable, "-m", "pip", "install", package], stdout=subprocess.DEVNULL)

install_package("python-docx")

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

# ── Màu sắc chuẩn ───────────────────────────────────────────────────────────
BLUE_DARK  = RGBColor(0x1F, 0x4E, 0x79)   # #1F4E79 — màu chính
BLUE_MID   = RGBColor(0x2E, 0x75, 0xB6)   # #2E75B6 — section header bg
WHITE      = RGBColor(0xFF, 0xFF, 0xFF)
GRAY_LIGHT = RGBColor(0xD6, 0xE4, 0xF0)   # header table row bg
GRAY_TEXT  = RGBColor(0x66, 0x66, 0x66)
FONT_NAME  = "Calibri"

# ── Helpers ──────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color: str):
    """Set background color of a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    # Remove existing shd
    for shd in tcPr.findall(qn('w:shd')):
        tcPr.remove(shd)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_borders(table, border_color="BDD7EE"):
    """Thin border on table cells."""
    for row in table.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            for side in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                border = OxmlElement(f'w:{side}')
                border.set(qn('w:val'), 'single')
                border.set(qn('w:sz'), '4')
                border.set(qn('w:color'), border_color)
                tcBorders.append(border)
            # Remove old
            for old in tcPr.findall(qn('w:tcBorders')):
                tcPr.remove(old)
            tcPr.append(tcBorders)

def add_run(para, text, bold=False, italic=False, size_pt=10, color=None):
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = FONT_NAME
    run.font.size = Pt(size_pt)
    if color:
        run.font.color.rgb = color
    return run

def set_para_spacing(para, before=0, after=0):
    pf = para.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)

# ── Block: Tiêu đề chính (bảng 1 cột, nền BLUE_DARK) ───────────────────────

def add_main_header(doc, project_name: str, subtitle: str = ""):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_bg(cell, "1F4E79")
    para = cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_spacing(para, 6, 4)
    add_run(para, "INTERVIEW TEMPLATE\n", bold=True, size_pt=18, color=WHITE)
    add_run(para, f"Dự án: {project_name}", bold=True, size_pt=13, color=WHITE)
    if subtitle:
        para.add_run("\n")
        add_run(para, subtitle, bold=False, size_pt=11, color=RGBColor(0xBD, 0xD7, 0xEE))

# ── Block: Bảng thông tin phỏng vấn ─────────────────────────────────────────

def add_info_table(doc, project_name: str, customer: str, created_date: str):
    rows_data = [
        ("Dự án",                project_name),
        ("Khách hàng",           customer),
        ("Ngày phỏng vấn",       ""),
        ("Người phỏng vấn (BA)", ""),
        ("Người được phỏng vấn", ""),
        ("Bộ phận / Vai trò",    ""),
    ]
    table = doc.add_table(rows=len(rows_data), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Column widths
    for row in table.rows:
        row.cells[0].width = Cm(5)
        row.cells[1].width = Cm(11)

    for i, (label, value) in enumerate(rows_data):
        bg = "D6E4F0" if i % 2 == 0 else "EBF3FB"
        set_cell_bg(table.cell(i, 0), bg)
        set_cell_bg(table.cell(i, 1), "FFFFFF")
        p0 = table.cell(i, 0).paragraphs[0]
        p1 = table.cell(i, 1).paragraphs[0]
        set_para_spacing(p0, 3, 3)
        set_para_spacing(p1, 3, 3)
        add_run(p0, label, bold=True, size_pt=10, color=BLUE_DARK)
        add_run(p1, value, bold=False, size_pt=10)

    set_cell_borders(table, "BDD7EE")
    doc.add_paragraph()  # spacer

# ── Block: Mục tiêu buổi phỏng vấn ─────────────────────────────────────────

def add_objectives(doc, objectives: list):
    p = doc.add_paragraph()
    set_para_spacing(p, 6, 2)
    add_run(p, "Mục tiêu buổi phỏng vấn:", bold=True, size_pt=11, color=BLUE_DARK)

    table = doc.add_table(rows=len(objectives), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, obj in enumerate(objectives):
        set_cell_bg(table.cell(i, 0), "1F4E79" if i == 0 else "2E75B6")
        set_cell_bg(table.cell(i, 1), "EBF3FB" if i % 2 == 0 else "FFFFFF")
        p0 = table.cell(i, 0).paragraphs[0]
        p1 = table.cell(i, 1).paragraphs[0]
        p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_para_spacing(p0, 3, 3)
        set_para_spacing(p1, 3, 3)
        add_run(p0, "✓", bold=True, size_pt=10, color=WHITE)
        add_run(p1, obj, bold=False, size_pt=10)
    set_cell_borders(table, "BDD7EE")
    doc.add_paragraph()

# ── Block: Section Header (stakeholder) ─────────────────────────────────────

def add_section_header(doc, emoji: str, title: str, subtitle: str = ""):
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    set_cell_bg(cell, "2E75B6")
    para = cell.paragraphs[0]
    set_para_spacing(para, 5, 5)
    add_run(para, f"{emoji}  {title}", bold=True, size_pt=14, color=WHITE)
    if subtitle:
        para.add_run("\n")
        add_run(para, subtitle, bold=False, size_pt=9, color=RGBColor(0xBD, 0xD7, 0xEE))
    doc.add_paragraph()

# ── Block: Sub-section label ─────────────────────────────────────────────────

def add_subsection_label(doc, label: str):
    p = doc.add_paragraph()
    set_para_spacing(p, 8, 2)
    add_run(p, label, bold=True, size_pt=11, color=BLUE_DARK)

# ── Block: Question table (3 cột: # | Câu hỏi | Ghi chú) ───────────────────

def add_question_table(doc, questions: list):
    """questions = list of (id, question_text)"""
    if not questions:
        return
    n = len(questions)
    table = doc.add_table(rows=n + 1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    headers = ["#", "Câu hỏi", "Ghi chú / Câu trả lời"]
    widths  = [Cm(1), Cm(9), Cm(6)]
    for j, (h, w) in enumerate(zip(headers, widths)):
        cell = table.cell(0, j)
        cell.width = w
        set_cell_bg(cell, "1F4E79")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_para_spacing(p, 3, 3)
        add_run(p, h, bold=True, size_pt=10, color=WHITE)

    # Data rows
    for i, (qid, qtext) in enumerate(questions):
        bg = "EBF3FB" if i % 2 == 0 else "FFFFFF"
        row = table.rows[i + 1]
        row.cells[0].width = widths[0]
        row.cells[1].width = widths[1]
        row.cells[2].width = widths[2]

        set_cell_bg(row.cells[0], bg)
        set_cell_bg(row.cells[1], bg)
        set_cell_bg(row.cells[2], "FFFFFF")

        p0 = row.cells[0].paragraphs[0]
        p1 = row.cells[1].paragraphs[0]
        p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_para_spacing(p0, 4, 4)
        set_para_spacing(p1, 4, 4)
        set_para_spacing(row.cells[2].paragraphs[0], 4, 4)

        add_run(p0, str(i + 1), bold=False, size_pt=10, color=BLUE_DARK)
        add_run(p1, qtext, bold=False, size_pt=10)

    set_cell_borders(table, "BDD7EE")
    doc.add_paragraph()

# ── Block: Tổng kết + Footer ─────────────────────────────────────────────────

def add_summary_block(doc):
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    set_cell_bg(cell, "EBF3FB")
    para = cell.paragraphs[0]
    set_para_spacing(para, 5, 5)
    add_run(para, "📝 TỔNG KẾT BUỔI PHỎNG VẤN", bold=True, size_pt=12, color=BLUE_DARK)
    para.add_run("\n")
    for label in ["Pain point chính:", "Quyết định / Action items:", "Artifacts cần xin:"]:
        add_run(para, f"\n• {label}  _______________________________________________", bold=False, size_pt=10)
    doc.add_paragraph()

def add_footer_line(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_spacing(p, 10, 0)
    add_run(p, "Người phỏng vấn ký xác nhận: ________________________     Ngày: _______________",
            bold=False, size_pt=9, color=GRAY_TEXT)

# ── Parser: đọc file MD ───────────────────────────────────────────────────────

def extract_project_info(lines):
    """Trích xuất tên dự án và ngày tạo từ header file."""
    project = "Dự án"
    date = ""
    customer = ""
    for line in lines[:10]:
        if line.startswith("# "):
            project = line[2:].strip()
        m = re.search(r'\*\*Ngày tạo:\*\*\s*(\S+)', line)
        if m:
            date = m.group(1)
        m2 = re.search(r'\*\*Đối tượng KH:\*\*\s*(.+)', line)
        if m2:
            customer = m2.group(1).strip()
    return project, date, customer

STAKEHOLDER_MAP = {
    "sales": ("📊", "PHỎNG VẤN SALES / KINH DOANH"),
    "kinh doanh": ("📊", "PHỎNG VẤN SALES / KINH DOANH"),
    "kho": ("🏭", "PHỎNG VẤN NHÂN VIÊN KHO / VẬN HÀNH"),
    "logistics": ("🏭", "PHỎNG VẤN NHÂN VIÊN KHO / VẬN HÀNH"),
    "kỹ thuật": ("🔧", "PHỎNG VẤN KỸ THUẬT VIÊN / TRIỂN KHAI"),
    "kế toán": ("💰", "PHỎNG VẤN KẾ TOÁN / TÀI CHÍNH"),
    "tài chính": ("💰", "PHỎNG VẤN KẾ TOÁN / TÀI CHÍNH"),
    "giám đốc": ("👔", "PHỎNG VẤN GIÁM ĐỐC / C-LEVEL"),
    "quản lý": ("👔", "PHỎNG VẤN QUẢN LÝ / DIRECTOR"),
    "it": ("💻", "PHỎNG VẤN IT / TECHNICAL TEAM"),
    "technical": ("💻", "PHỎNG VẤN IT / TECHNICAL TEAM"),
}

SECTION_LABEL_MAP = {
    "1.1": "1.1 — Vấn đề & Mục đích",
    "1.2": "1.2 — Hệ thống hiện tại",
    "1.2b": "1.2B — Hệ thống hiện tại (Chuyên sâu)",
    "1.3a": "1.3A — Xác nhận Stakeholder",
    "1.3b": "1.3B — Vòng đời thực thể",
    "1.3c": "1.3C — Công việc chi tiết từng Stakeholder",
    "1.4": "1.4 — Phân tích Tác động",
    "1.5": "1.5 — Thu thập Artifacts",
    "2": "2 — Đối thủ & Thị trường",
    "3": "3 — Nhu cầu tương lai & MVP",
    "4": "4 — User Story",
    "5": "5 — Tích hợp & Đồng bộ",
    "6": "6 — Yêu cầu Kỹ thuật",
    "7": "7 — Báo cáo & Dashboard",
    "8": "8 — Ngân sách & Timeline",
    "9": "9 — Câu hỏi Bổ sung",
}

def get_section_label(section_id: str) -> str:
    sid = section_id.lower().strip()
    for key, val in SECTION_LABEL_MAP.items():
        if sid.startswith(key):
            return val
    return section_id

def detect_stakeholder(heading: str):
    h = heading.lower()
    for key, (emoji, title) in STAKEHOLDER_MAP.items():
        if key in h:
            return emoji, title
    return "🗣️", heading.replace("👤", "").strip()

def clean_question_text(text: str) -> str:
    """Làm sạch text câu hỏi từ bảng MD."""
    # Bỏ | ký tự thừa
    text = text.strip().strip('|').strip()
    return text

def parse_md_table(lines, start_idx):
    """Parse một bảng MD, trả về (headers, rows, end_idx)."""
    rows = []
    i = start_idx
    while i < len(lines):
        line = lines[i].strip()
        if not line.startswith('|'):
            break
        if re.match(r'^\|[\s\-:|]+\|', line):
            i += 1
            continue
        cells = [c.strip() for c in line.split('|')[1:-1]]
        rows.append(cells)
        i += 1
    return rows, i

def extract_questions_from_table(rows, id_col=0, question_cols=None):
    """
    Trích xuất câu hỏi từ table rows.
    question_cols: list các index cột chứa câu hỏi
    """
    questions = []
    if question_cols is None:
        question_cols = [2] if len(rows) > 0 and len(rows[0]) >= 3 else [1]

    for row in rows:
        if not row:
            continue
        # Lấy ID
        qid = row[id_col].strip() if len(row) > id_col else ""
        if qid.upper() in ("ID", "#", "MỤC") or (not qid and len(row) > 1):
            continue
        # Gộp tất cả các cột câu hỏi có nội dung
        parts = []
        for ci in question_cols:
            if ci < len(row) and row[ci].strip():
                parts.append(clean_question_text(row[ci]))
        if parts:
            questions.append((qid, " | ".join(parts)))
    return questions

def md_to_interview_docx(md_file: str, docx_file: str):
    if not os.path.exists(md_file):
        print(f"✗ File không tìm thấy: {md_file}")
        return False

    with open(md_file, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    lines_stripped = [l.rstrip('\n').rstrip('\r') for l in lines]

    # ── Thông tin dự án
    project_name, created_date, customer = extract_project_info(lines_stripped)
    if not customer:
        customer = "—"

    # ── Tạo document
    doc = Document()
    # Page margins
    for section in doc.sections:
        section.top_margin    = Cm(1.8)
        section.bottom_margin = Cm(1.8)
        section.left_margin   = Cm(2.0)
        section.right_margin  = Cm(2.0)

    # ── Header chính
    add_main_header(doc, project_name, f"Ngày tạo: {created_date}")
    doc.add_paragraph()

    # ── Bảng thông tin
    add_info_table(doc, project_name, customer, created_date)

    # ── Mục tiêu mặc định
    objectives = [
        "Hiểu hiện trạng hệ thống & quy trình vận hành",
        "Xác định pain point cụ thể theo từng nghiệp vụ",
        "Xác định scope và ưu tiên tính năng cho MVP",
        "Thu thập thông tin để viết User Story / Use Case",
    ]
    add_objectives(doc, objectives)

    # ── Parse nội dung
    i = 0
    current_section = None
    current_stakeholder = None
    pending_questions = []   # buffer khi chưa có stakeholder

    def flush_questions(label=None):
        """Flush question buffer ra doc."""
        if not pending_questions:
            return
        if label:
            add_subsection_label(doc, label)
        add_question_table(doc, pending_questions)
        pending_questions.clear()

    section_order = []  # track thứ tự section

    while i < len(lines_stripped):
        line = lines_stripped[i]

        # ── Heading cấp 2: PHẦN ...
        if re.match(r'^##\s+PHẦN', line, re.IGNORECASE):
            flush_questions(current_section)
            current_section = line.lstrip('#').strip()
            # Detect sub-section
            m = re.search(r'PHẦN\s+([\d\.AB]+)', current_section, re.IGNORECASE)
            if m:
                current_section = get_section_label(m.group(1))
            current_stakeholder = None
            i += 1
            continue

        # ── Heading cấp 3: ### Nhóm... hoặc stakeholder
        if line.startswith('### '):
            flush_questions(current_section if current_stakeholder is None else current_section)
            heading = line.lstrip('#').strip()
            # Detect stakeholder
            emoji, title = detect_stakeholder(heading)
            current_stakeholder = title
            add_section_header(doc, emoji, title, current_section or "")
            current_section = heading  # update label
            i += 1
            continue

        # ── Heading cấp 4: #### hoặc 👤
        if line.startswith('#### ') or (line.startswith('### ') and '👤' in line):
            flush_questions(current_section)
            heading = re.sub(r'^#+', '', line).replace('👤', '').strip()
            emoji, title = detect_stakeholder(heading)
            add_section_header(doc, emoji, title, "")
            current_section = title
            i += 1
            continue

        # ── Table row
        if line.strip().startswith('|'):
            rows, new_i = parse_md_table(lines_stripped, i)
            if len(rows) > 1:  # skip header-only
                # Xác định các cột câu hỏi
                header = [h.lower() for h in rows[0]] if rows else []
                q_cols = []
                
                # Định nghĩa các nhóm từ khóa
                CONTENT_KEYWORDS = ['câu hỏi', 'question', 'mục tiêu', 'mô tả', 'trạng thái', 'stakeholder', 'vai trò', 'nội dung', 'yêu cầu', 'mục', 'đối thủ', 'xu hướng']
                STATUS_KEYWORDS = ['status', 'xác nhận', 'loại', 'ngày', 'ghi chú']
                
                # Ưu tiên các cột chứa từ khóa CONTENT nhưng KHÔNG chứa từ khóa STATUS
                for ci, h in enumerate(header):
                    is_content = any(kw in h for kw in CONTENT_KEYWORDS)
                    is_status = any(kw in h for kw in STATUS_KEYWORDS)
                    
                    if is_content and not is_status:
                        q_cols.append(ci)
                
                # Nếu không tìm thấy cột Content theo keyword, dùng fallback logic
                if not q_cols:
                    # Lấy tất cả các cột trừ cột ID (0) và các cột Status/Metadata
                    for ci, h in enumerate(header):
                        if ci == 0: continue # thường là ID
                        if not any(kw in h for kw in STATUS_KEYWORDS):
                            q_cols.append(ci)
                
                # Nếu vẫn trống, lấy cột 1 hoặc 2 làm fallback cuối cùng
                if not q_cols:
                    q_cols = [2] if len(header) >= 3 else ([1] if len(header) >= 2 else [0])

                qs = extract_questions_from_table(rows[1:], id_col=0, question_cols=q_cols)
                pending_questions.extend(qs)
            i = new_i
            continue

        i += 1

    # Flush cuối
    flush_questions(current_section)

    # ── Tổng kết & Footer
    doc.add_page_break()
    add_summary_block(doc)
    add_footer_line(doc)

    # ── Lưu file
    try:
        doc.save(docx_file)
        print(f"✓ Saved: {docx_file}")
        return True
    except Exception as e:
        print(f"✗ Failed to save: {e}")
        print("TIP: Đóng file DOCX trong Word nếu đang mở.")
        return False


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: py md_qna_to_interview.py <input.md> [output.docx]")
        sys.exit(1)

    md_input = sys.argv[1]

    if len(sys.argv) >= 3:
        docx_output = sys.argv[2]
    else:
        # Mặc định: cùng thư mục với file input
        base = os.path.basename(md_input)
        name_no_ext = os.path.splitext(base)[0]
        # Chuyển questions_tracking_YYYYMMDD_slug → Interview_slug_YYYYMMDD
        m = re.match(r'questions_tracking_(\d{8})_(.+)', name_no_ext)
        if m:
            docx_name = f"Interview_{m.group(2)}_{m.group(1)}.docx"
        else:
            docx_name = f"Interview_{name_no_ext}.docx"
        docx_output = os.path.join(os.path.dirname(md_input), docx_name)

    success = md_to_interview_docx(md_input, docx_output)
    sys.exit(0 if success else 1)
