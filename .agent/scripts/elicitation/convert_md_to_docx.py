#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Script chuyển đổi file Markdown sang Word (.docx)
Hỗ trợ các file: khơi gợi yêu cầu, summary, tracking questions
"""

import sys
import os
import re
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("Đang cài đặt thư viện python-docx...")
    os.system("pip install python-docx")
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

def parse_markdown_table(line):
    """Parse markdown table row"""
    cells = [cell.strip() for cell in line.split('|') if cell.strip()]
    return cells

def is_table_row(line):
    """Kiểm tra xem dòng có phải là table row không"""
    return '|' in line and not line.strip().startswith('|---')

def parse_markdown(md_file):
    """Parse markdown file thành các phần tử"""
    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    lines = content.split('\n')
    elements = []
    current_table = None
    current_code_block = None
    current_list = None
    
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        
        # Code block
        if stripped.startswith('```'):
            if current_code_block is None:
                current_code_block = [line]
            else:
                current_code_block.append(line)
                elements.append(('code', '\n'.join(current_code_block)))
                current_code_block = None
            i += 1
            continue
        
        if current_code_block is not None:
            current_code_block.append(line)
            i += 1
            continue
        
        # Headers
        if stripped.startswith('#'):
            # Đóng table trước nếu đang mở
            if current_table:
                elements.append(('table', current_table))
                current_table = None
            
            level = len(stripped) - len(stripped.lstrip('#'))
            text = stripped.lstrip('#').strip()
            elements.append(('header', level, text))
            i += 1
            continue
        
        # Table
        if is_table_row(line):
            if current_table is None:
                current_table = {'headers': None, 'rows': []}
            
            cells = parse_markdown_table(line)
            if current_table['headers'] is None:
                # Check if next line is separator
                if i + 1 < len(lines) and '|---' in lines[i + 1]:
                    current_table['headers'] = cells
                    i += 2  # Skip header and separator
                    continue
            
            if current_table['headers'] is not None:
                current_table['rows'].append(cells)
        else:
            # Đóng table nếu đang mở
            if current_table:
                elements.append(('table', current_table))
                current_table = None
            
            # List items
            if stripped.startswith('- ') or stripped.startswith('* '):
                if current_list is None:
                    current_list = []
                text = stripped[2:].strip()
                # Parse formatting
                current_list.append(text)
            elif stripped.startswith(('  ', '\t')) and (stripped.lstrip().startswith('- ') or stripped.lstrip().startswith('* ')):
                if current_list is None:
                    current_list = []
                text = stripped.lstrip()[2:].strip()
                current_list.append(text)
            elif re.match(r'^\d+\.\s', stripped):
                if current_list is None:
                    current_list = []
                text = re.sub(r'^\d+\.\s', '', stripped)
                current_list.append(text)
            else:
                # Đóng list nếu đang mở
                if current_list is not None:
                    elements.append(('list', current_list))
                    current_list = None
                
                # Paragraph
                if stripped:
                    elements.append(('paragraph', stripped))
                else:
                    elements.append(('empty',))
        
        i += 1
    
    # Đóng các phần tử còn lại
    if current_table:
        elements.append(('table', current_table))
    if current_list is not None:
        elements.append(('list', current_list))
    
    return elements

def format_text_with_style(text):
    """Parse markdown formatting trong text"""
    # Bold: **text** hoặc __text__
    text = re.sub(r'\*\*(.*?)\*\*', r'<b>\1</b>', text)
    text = re.sub(r'__(.*?)__', r'<b>\1</b>', text)
    
    # Italic: *text* hoặc _text_
    text = re.sub(r'(?<!\*)\*(?!\*)(.*?)(?<!\*)\*(?!\*)', r'<i>\1</i>', text)
    text = re.sub(r'(?<!_)_(?!_)(.*?)(?<!_)_(?!_)', r'<i>\1</i>', text)
    
    # Code: `text`
    text = re.sub(r'`([^`]+)`', r'<code>\1</code>', text)
    
    # Links: [text](url)
    text = re.sub(r'\[([^\]]+)\]\(([^\)]+)\)', r'\1', text)
    
    return text

def apply_formatting(paragraph, text):
    """Áp dụng formatting cho paragraph"""
    parts = re.split(r'(<b>.*?</b>|<i>.*?</i>|<code>.*?</code>)', text)
    for part in parts:
        if part.startswith('<b>') and part.endswith('</b>'):
            run_bold = paragraph.add_run(part[3:-4])
            run_bold.bold = True
        elif part.startswith('<i>') and part.endswith('</i>'):
            run_italic = paragraph.add_run(part[3:-4])
            run_italic.italic = True
        elif part.startswith('<code>') and part.endswith('</code>'):
            run_code = paragraph.add_run(part[6:-7])
            run_code.font.name = 'Courier New'
        elif part:
            paragraph.add_run(part)

def convert_md_to_docx(md_file, docx_file):
    """Chuyển đổi Markdown sang Word"""
    
    try:
        print(f"Dang doc file: {md_file}")
    except:
        pass
    elements = parse_markdown(md_file)
    
    # Tạo Word document
    doc = Document()
    
    # Cấu hình font mặc định
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    # Process elements
    for element in elements:
        if element[0] == 'header':
            level = element[1]
            text = format_text_with_style(element[2])
            
            heading = doc.add_heading(level=min(level, 9))
            # Clear default run
            heading.clear()
            apply_formatting(heading, text)
            
        elif element[0] == 'paragraph':
            text = format_text_with_style(element[1])
            para = doc.add_paragraph()
            apply_formatting(para, text)
            
        elif element[0] == 'list':
            for item in element[1]:
                text = format_text_with_style(item)
                para = doc.add_paragraph(style='List Bullet')
                if para.runs:
                    para.clear()
                apply_formatting(para, text)
                
        elif element[0] == 'table':
            table_data = element[1]
            if table_data['headers'] and table_data['rows']:
                # Tạo table
                num_cols = len(table_data['headers'])
                num_rows = len(table_data['rows']) + 1  # +1 for header
                table = doc.add_table(rows=num_rows, cols=num_cols)
                table.style = 'Light Grid Accent 1'
                
                # Header
                header_cells = table.rows[0].cells
                for i, header in enumerate(table_data['headers']):
                    header_cells[i].text = header.replace('**', '').strip()
                    # Bold header
                    for paragraph in header_cells[i].paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
                
                # Rows
                for row_idx, row_data in enumerate(table_data['rows']):
                    row_cells = table.rows[row_idx + 1].cells
                    for col_idx in range(num_cols):
                        if col_idx < len(row_data):
                            cell_text = row_data[col_idx].replace('**', '').replace('✅', 'Đã làm rõ').replace('⬜', 'Chưa làm rõ')
                            row_cells[col_idx].text = cell_text.strip()
                        else:
                            row_cells[col_idx].text = ''
        
        elif element[0] == 'code':
            # Code block - giữ nguyên format
            code_text = element[1]
            para = doc.add_paragraph()
            run = para.add_run(code_text)
            run.font.name = 'Courier New'
            run.font.size = Pt(10)
        
        elif element[0] == 'empty':
            # Empty line
            doc.add_paragraph()
    
    # Lưu file
    doc.save(docx_file)
    print(f"SUCCESS: Da tao file Word: {docx_file}")

def main():
    if len(sys.argv) < 2:
        print("Usage: python convert_md_to_docx.py <markdown_file> [output_file]")
        print("Example: python convert_md_to_docx.py file.md")
        sys.exit(1)
    
    md_file_arg = sys.argv[1]
    
    # Xử lý đường dẫn
    md_file_path = Path(md_file_arg)
    if not md_file_path.is_absolute():
        md_file = str(md_file_path.resolve())
    else:
        md_file = str(md_file_path)
    
    if len(sys.argv) >= 3:
        docx_file_arg = sys.argv[2]
        docx_file_path = Path(docx_file_arg)
        if not docx_file_path.is_absolute():
            docx_file = str(docx_file_path.resolve())
        else:
            docx_file = str(docx_file_path)
    else:
        # Tạo output file cùng thư mục với input file
        md_file_path_obj = Path(md_file)
        docx_file = str(md_file_path_obj.with_suffix('.docx'))
    
    if not os.path.exists(md_file):
        print(f"ERROR: Khong tim thay file: {md_file}")
        print(f"   Duong dan hien tai: {os.getcwd()}")
        sys.exit(1)
    
    convert_md_to_docx(md_file, docx_file)

if __name__ == '__main__':
    main()
