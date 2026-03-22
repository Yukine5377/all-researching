#!/usr/bin/env python3
"""
Convert Markdown to Word (.docx) for Meeting Minutes
Safe Version: No external style dependencies
Handles file locks and provides better feedback.
Author: BA Team
Version: 1.0.4
"""

import sys
import os
import re
import subprocess

def install_package(package):
    try:
        __import__(package.replace('-', '_'))
    except ImportError:
        print(f"Installing package '{package}'...")
        try:
            subprocess.check_call([sys.executable, "-m", "pip", "install", package])
        except Exception as e:
            print(f"✗ Failed to install '{package}': {e}")

install_package("python-docx")

from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_hyperlink(paragraph, text, url):
    try:
        part = paragraph.part
        r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
        hyperlink = OxmlElement('w:hyperlink')
        hyperlink.set(qn('r:id'), r_id)
        new_run = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')
        c = OxmlElement('w:color')
        c.set(qn('w:val'), '0000FF')
        rPr.append(c)
        u = OxmlElement('w:u')
        u.set(qn('w:val'), 'single')
        rPr.append(u)
        new_run.append(rPr)
        t = OxmlElement('w:t')
        t.set(qn('xml:space'), 'preserve')
        t.text = text
        new_run.append(t)
        hyperlink.append(new_run)
        paragraph._p.append(hyperlink)
        return hyperlink
    except:
        return paragraph.add_run(f"{text} ({url})")

def parse_inline_formatting(text, para):
    link_pattern = r'\[([^\]]+)\]\(([^\)]+)\)'
    last_end = 0
    for match in re.finditer(link_pattern, text):
        if match.start() > last_end:
            process_bold_italic(text[last_end:match.start()], para)
        add_hyperlink(para, match.group(1), match.group(2))
        last_end = match.end()
    if last_end < len(text):
        process_bold_italic(text[last_end:], para)

def process_bold_italic(text, para):
    # Fixed regex to correctly capture groups
    pattern = r'(\*\*.*?\*\*)|(\*.*?\*)'
    last_end = 0
    for match in re.finditer(pattern, text):
        if match.start() > last_end:
            para.add_run(text[last_end:match.start()])
        if match.group(1):
            para.add_run(match.group(1)[2:-2]).bold = True
        elif match.group(2):
            para.add_run(match.group(2)[1:-1]).italic = True
        last_end = match.end()
    if last_end < len(text):
        para.add_run(text[last_end:])

def markdown_to_docx(md_file, docx_file):
    if not os.path.exists(md_file):
        print(f"Error: {md_file} not found")
        return False
    
    with open(md_file, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    doc = Document()
    doc.styles['Normal'].font.name = 'Times New Roman'
    doc.styles['Normal'].font.size = Pt(12)

    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        if not line.strip(): i += 1; continue

        if line.startswith('#'):
            level = len(line.split()[0])
            p = doc.add_heading('', level=min(level, 4))
            parse_inline_formatting(line.lstrip('#').strip(), p)
            i += 1; continue

        if line.strip().startswith('|'):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i]); i += 1
            if len(table_lines) >= 2:
                header = [c.strip() for c in table_lines[0].split('|')[1:-1]]
                table = doc.add_table(rows=1, cols=len(header))
                table.style = 'Table Grid'
                for idx, text in enumerate(header):
                    table.rows[0].cells[idx].paragraphs[0].add_run(text).bold = True
                start_row = 2 if '-' in table_lines[1] else 1
                for row_line in table_lines[start_row:]:
                    row_data = [c.strip() for c in row_line.split('|')[1:-1]]
                    row = table.add_row()
                    for idx, text in enumerate(row_data):
                        if idx < len(row.cells): parse_inline_formatting(text, row.cells[idx].paragraphs[0])
            continue

        if line.strip().startswith(('- ', '* ')):
            p = doc.add_paragraph(); p.add_run('•\t')
            parse_inline_formatting(line.strip()[2:].strip(), p)
            i += 1; continue

        p = doc.add_paragraph()
        parse_inline_formatting(line.strip(), p)
        i += 1

    try:
        doc.save(docx_file)
        print(f"✓ Saved: {docx_file}")
        return True
    except Exception as e:
        print(f"✗ Failed to save to {docx_file}: {e}")
        print("TIP: Make sure the file is not open in Microsoft Word.")
        return False

if __name__ == "__main__":
    if len(sys.argv) < 3:
        print("Usage: py md_to_docx.py <input.md> <output.docx>")
        sys.exit(1)
    if markdown_to_docx(sys.argv[1], sys.argv[2]):
        sys.exit(0)
    else:
        sys.exit(1)
