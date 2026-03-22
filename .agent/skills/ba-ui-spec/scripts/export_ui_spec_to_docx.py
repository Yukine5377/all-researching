#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Script xuất UI Specification ra file Word (.docx)
Hỗ trợ format bảng chuyên nghiệp cho tài liệu UI Spec
"""

import sys
import os
import re
from pathlib import Path
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("Đang cài đặt thư viện python-docx...")
    os.system("pip install python-docx")
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

def set_cell_width(cell, width_inches):
    """Set width cho cell trong table"""
    cell_width = cell._element
    cell_width.set(qn('w:w'), str(int(width_inches * 1440)))
    cell_width.set(qn('w:type'), 'dxa')

def parse_text_table(content):
    """Parse bảng dạng text (không phải markdown)"""
    lines = content.strip().split('\n')
    table_data = []
    
    for line in lines:
        # Bỏ qua dòng trống
        if not line.strip():
            continue
        
        # Kiểm tra nếu là header (có chứa các từ khóa như STT, Tên control, v.v.)
        if 'STT' in line or 'Tên control' in line or 'Loại control' in line:
            # Parse header
            if '|' in line:
                cells = [cell.strip() for cell in line.split('|') if cell.strip()]
            else:
                # Nếu không có |, thử tách bằng tab hoặc nhiều khoảng trắng
                cells = re.split(r'\t+|\s{2,}', line.strip())
            if cells:
                table_data.append({'type': 'header', 'cells': cells})
            continue
        
        # Parse rows - kiểm tra nếu có số thứ tự ở đầu
        if re.match(r'^\d+', line.strip()):
            if '|' in line:
                cells = [cell.strip() for cell in line.split('|') if cell.strip()]
            else:
                cells = re.split(r'\t+|\s{2,}', line.strip())
            if len(cells) >= 3:  # Ít nhất phải có 3 cột
                table_data.append({'type': 'row', 'cells': cells})
    
    return table_data

def parse_markdown_table(content):
    """Parse bảng markdown format"""
    lines = content.strip().split('\n')
    headers = None
    rows = []
    in_table = False
    
    for line in lines:
        stripped = line.strip()
        
        # Bỏ qua separator line
        if stripped.startswith('|---') or re.match(r'^\|[-:|\s]+\|', stripped):
            continue
        
        # Parse header
        if '|' in line and headers is None:
            headers = [cell.strip() for cell in line.split('|') if cell.strip()]
            in_table = True
            continue
        
        # Parse rows
        if in_table and '|' in line:
            cells = [cell.strip() for cell in line.split('|') if cell.strip()]
            if cells and len(cells) == len(headers):
                rows.append(cells)
        elif in_table and not stripped:
            in_table = False
    
    if headers and rows:
        return {'headers': headers, 'rows': rows}
    return None

def create_ui_spec_docx(ui_spec_file, output_file=None):
    """Tạo file Word từ UI Spec"""
    
    # Đọc file
    with open(ui_spec_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Tạo Word document
    doc = Document()
    
    # Cấu hình font mặc định
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    # Title
    title = doc.add_heading('ĐẶC TẢ GIAO DIỆN MÀN HÌNH (UI SPECIFICATION)', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Parse nội dung
    sections = content.split('\n\n')
    
    # Tìm thông tin màn hình
    screen_info = {}
    for line in content.split('\n'):
        line_stripped = line.strip()
        # Parse các format khác nhau
        if re.match(r'^[Tt]ên chức năng\s*[:：]', line_stripped) or re.match(r'^Chức năng\s*[:：]', line_stripped):
            screen_info['function'] = re.split(r'[:：]', line_stripped, 1)[1].strip()
        elif re.match(r'^Actor\s*[:：]', line_stripped) or re.match(r'^Ai sử dụng\s*[:：]', line_stripped):
            screen_info['user'] = re.split(r'[:：]', line_stripped, 1)[1].strip()
        elif re.match(r'^Mục đích\s*[:：]', line_stripped):
            screen_info['purpose'] = re.split(r'[:：]', line_stripped, 1)[1].strip()
        elif re.match(r'^Hệ thống\s*[:：]', line_stripped):
            screen_info['system'] = re.split(r'[:：]', line_stripped, 1)[1].strip()
    
    # 1. Thông tin chung (dạng bảng)
    doc.add_paragraph()
    info_heading = doc.add_heading('1. Thông tin chung', 1)
    
    # Tạo bảng thông tin chung (3 hàng, 2 cột)
    info_table = doc.add_table(rows=3, cols=2)
    info_table.style = 'Light Grid Accent 1'
    
    # Set column widths
    set_cell_width(info_table.rows[0].cells[0], 2.5)  # Cột 1: Tên trường
    set_cell_width(info_table.rows[0].cells[1], 4.5)  # Cột 2: Giá trị
    
    # Row 1: Tên chức năng
    info_table.rows[0].cells[0].text = 'Tên chức năng'
    info_table.rows[0].cells[1].text = screen_info.get('function', '')
    # Bold cột đầu tiên
    for paragraph in info_table.rows[0].cells[0].paragraphs:
        for run in paragraph.runs:
            run.bold = True
    
    # Row 2: Actor
    info_table.rows[1].cells[0].text = 'Actor'
    info_table.rows[1].cells[1].text = screen_info.get('user', '')
    # Bold cột đầu tiên
    for paragraph in info_table.rows[1].cells[0].paragraphs:
        for run in paragraph.runs:
            run.bold = True
    
    # Row 3: Mục đích
    info_table.rows[2].cells[0].text = 'Mục đích'
    info_table.rows[2].cells[1].text = screen_info.get('purpose', '')
    # Bold cột đầu tiên
    for paragraph in info_table.rows[2].cells[0].paragraphs:
        for run in paragraph.runs:
            run.bold = True
    
    # 2. Màn hình
    doc.add_paragraph()
    screen_heading = doc.add_heading('2. Màn hình', 1)
    
    # Tìm tên màn hình từ file name hoặc từ thông tin
    screen_name = screen_info.get('function', 'Màn hình')
    # Nếu có "Màn hình:" trong file thì lấy từ đó
    for line in content.split('\n'):
        if re.match(r'^Màn hình\s*[:：]', line.strip()):
            screen_name = re.split(r'[:：]', line.strip(), 1)[1].strip()
            break
    
    screen_para = doc.add_paragraph()
    screen_para.add_run(screen_name)
    
    # 3. Mô tả màn hình
    doc.add_paragraph()
    table_heading = doc.add_heading('3. Mô tả màn hình', 1)
    
    # Thử parse markdown table trước
    table_data = parse_markdown_table(content)
    
    if not table_data:
        # Nếu không phải markdown, thử parse text table
        table_rows = parse_text_table(content)
        if table_rows:
            # Tìm header
            header_row = None
            data_rows = []
            for row in table_rows:
                if row['type'] == 'header':
                    header_row = row['cells']
                elif row['type'] == 'row':
                    data_rows.append(row['cells'])
            
            if header_row:
                table_data = {'headers': header_row, 'rows': [r['cells'] for r in data_rows]}
    
    # Tạo bảng nếu có data
    if table_data and table_data.get('headers') and table_data.get('rows'):
        headers = table_data['headers']
        rows = table_data['rows']
        
        num_cols = len(headers)
        num_rows = len(rows) + 1  # +1 for header
        
        table = doc.add_table(rows=num_rows, cols=num_cols)
        table.style = 'Light Grid Accent 1'
        
        # Set column widths
        if num_cols == 7:  # Create/Update/Delete
            widths = [0.5, 2.0, 1.5, 0.8, 1.0, 4.0, 2.5]  # Inches
        elif num_cols == 3:  # Read
            widths = [2.0, 1.5, 6.5]  # Inches
        else:
            widths = [2.0] * num_cols
        
        # Header
        header_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            # Remove markdown formatting
            header_text = header.replace('**', '').replace('*', '').strip()
            header_cells[i].text = header_text
            
            # Bold header
            for paragraph in header_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.bold = True
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Set width
            if i < len(widths):
                set_cell_width(header_cells[i], widths[i])
        
        # Rows
        for row_idx, row_data in enumerate(rows):
            row_cells = table.rows[row_idx + 1].cells
            for col_idx in range(num_cols):
                if col_idx < len(row_data):
                    cell_text = row_data[col_idx]
                    # Remove markdown formatting
                    cell_text = cell_text.replace('**', '').replace('*', '').strip()
                    
                    # Handle multi-line và bullet points
                    lines = cell_text.split('\n')
                    has_newline = len(lines) > 1
                    
                    if has_newline:
                        # Xóa text mặc định của cell
                        row_cells[col_idx].text = ''
                        first_para = True
                        
                        for line in lines:
                            line = line.strip()
                            if not line:
                                continue
                            
                            # Kiểm tra nếu dòng bắt đầu bằng bullet (•, -, *) với khoảng trắng sau đó
                            bullet_match = re.match(r'^[•\-\*]\s+(.+)$', line)
                            
                            if bullet_match:
                                # Tạo paragraph với bullet style
                                para = row_cells[col_idx].add_paragraph()
                                para.style = 'List Bullet'
                                para.text = bullet_match.group(1).strip()
                                first_para = False
                            else:
                                # Tạo paragraph bình thường
                                if first_para:
                                    para = row_cells[col_idx].paragraphs[0]
                                    para.text = line
                                    first_para = False
                                else:
                                    para = row_cells[col_idx].add_paragraph()
                                    para.text = line
                    else:
                        # Không có xuống dòng, kiểm tra các pattern để tự động tạo bullet points
                        # Pattern 1: Có bullet marker trong text (•, -, *) với khoảng trắng
                        if re.search(r'[•\-\*]\s+', cell_text):
                            row_cells[col_idx].text = ''
                            first_para = True
                            
                            # Split theo bullet pattern
                            parts = re.split(r'([•\-\*]\s+)', cell_text)
                            
                            current_text = ''
                            for i, part in enumerate(parts):
                                if re.match(r'^[•\-\*]\s+$', part):
                                    # Bullet marker, xử lý text trước đó (nếu có)
                                    if current_text.strip():
                                        if first_para:
                                            para = row_cells[col_idx].paragraphs[0]
                                            para.text = current_text.strip()
                                            first_para = False
                                        else:
                                            para = row_cells[col_idx].add_paragraph()
                                            para.text = current_text.strip()
                                        current_text = ''
                                elif re.match(r'^[•\-\*]\s+', part):
                                    # Text bắt đầu bằng bullet
                                    if current_text.strip():
                                        if first_para:
                                            para = row_cells[col_idx].paragraphs[0]
                                            para.text = current_text.strip()
                                            first_para = False
                                        else:
                                            para = row_cells[col_idx].add_paragraph()
                                            para.text = current_text.strip()
                                        current_text = ''
                                    # Thêm bullet item
                                    bullet_text = re.sub(r'^[•\-\*]\s+', '', part).strip()
                                    if bullet_text:
                                        para = row_cells[col_idx].add_paragraph()
                                        para.style = 'List Bullet'
                                        para.text = bullet_text
                                        first_para = False
                                else:
                                    # Text bình thường
                                    current_text += part
                            
                            # Thêm text cuối cùng nếu có
                            if current_text.strip():
                                if first_para:
                                    para = row_cells[col_idx].paragraphs[0]
                                    para.text = current_text.strip()
                                else:
                                    para = row_cells[col_idx].add_paragraph()
                                    para.text = current_text.strip()
                        # Pattern 2: Có nhiều mục được liệt kê sau dấu hai chấm, phân tách bằng dấu phẩy
                        # Ví dụ: 'Mô tả chính: "A", "B", "C"' hoặc có/không có ngoặc đơn
                        elif ':' in cell_text and ',' in cell_text:
                            # Phát hiện có nhiều mục trong phần liệt kê
                            row_cells[col_idx].text = ''
                            first_para = True
                            
                            # Tách phần mô tả chính và phần liệt kê
                            # Tìm vị trí bắt đầu của phần liệt kê (thường có dấu hai chấm)
                            colon_pos = cell_text.find(':')
                            
                            if colon_pos > 0:
                                # Phần mô tả trước dấu hai chấm
                                main_desc = cell_text[:colon_pos + 1].strip()
                                para = row_cells[col_idx].paragraphs[0]
                                para.text = main_desc
                                first_para = False
                                
                                # Phần liệt kê sau dấu hai chấm
                                list_part = cell_text[colon_pos + 1:].strip()
                                
                                # Tách theo dấu phẩy, nhưng chỉ tách ở level ngoài cùng (không tách trong ngoặc đơn)
                                items = []
                                current_item = ''
                                paren_depth = 0
                                
                                for char in list_part:
                                    if char == '(':
                                        paren_depth += 1
                                        current_item += char
                                    elif char == ')':
                                        paren_depth -= 1
                                        current_item += char
                                    elif char == ',' and paren_depth == 0:
                                        # Dấu phẩy ở level ngoài cùng
                                        if current_item.strip():
                                            items.append(current_item.strip())
                                        current_item = ''
                                    else:
                                        current_item += char
                                
                                # Thêm item cuối cùng
                                if current_item.strip():
                                    items.append(current_item.strip())
                                
                                # Chỉ tạo bullet nếu có ít nhất 2 mục
                                if len(items) >= 2:
                                    for item in items:
                                        para = row_cells[col_idx].add_paragraph()
                                        para.style = 'List Bullet'
                                        para.text = item
                                else:
                                    # Nếu không đủ điều kiện coi là list, hiển thị bình thường
                                    row_cells[col_idx].text = cell_text
                            else:
                                # Không có dấu hai chấm, hiển thị bình thường
                                row_cells[col_idx].text = cell_text
                        else:
                            # Không có pattern đặc biệt, hiển thị bình thường
                            row_cells[col_idx].text = cell_text
                    
                    # Center align cho cột STT, Require, MaxLength
                    if col_idx in [0, 3, 4] and num_cols == 7:
                        for paragraph in row_cells[col_idx].paragraphs:
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    elif col_idx == 0 and num_cols == 3:
                        for paragraph in row_cells[col_idx].paragraphs:
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                else:
                    row_cells[col_idx].text = ''
    else:
        # Nếu không có bảng, thêm thông báo
        doc.add_paragraph()
        doc.add_paragraph("Không tìm thấy bảng đặc tả trong file.")
    
    # Lưu file
    if output_file is None:
        input_path = Path(ui_spec_file)
        output_file = str(input_path.with_suffix('.docx'))
    
    doc.save(output_file)
    try:
        print(f"SUCCESS: Da tao file Word: {output_file}")
    except UnicodeEncodeError:
        print(f"SUCCESS: Created Word file: {output_file}")
    return output_file

def main():
    if len(sys.argv) < 2:
        print("Usage: python export_ui_spec_to_docx.py <ui_spec_file> [output_file]")
        print("Example: python export_ui_spec_to_docx.py ui_spec.md")
        sys.exit(1)
    
    ui_spec_file_arg = sys.argv[1]
    
    # Xử lý đường dẫn
    ui_spec_file_path = Path(ui_spec_file_arg)
    if not ui_spec_file_path.is_absolute():
        ui_spec_file = str(ui_spec_file_path.resolve())
    else:
        ui_spec_file = str(ui_spec_file_path)
    
    if len(sys.argv) >= 3:
        output_file_arg = sys.argv[2]
        output_file_path = Path(output_file_arg)
        if not output_file_path.is_absolute():
            output_file = str(output_file_path.resolve())
        else:
            output_file = str(output_file_path)
    else:
        output_file = None
    
    if not os.path.exists(ui_spec_file):
        try:
            print(f"ERROR: Khong tim thay file: {ui_spec_file}")
            print(f"   Duong dan hien tai: {os.getcwd()}")
        except UnicodeEncodeError:
            print(f"ERROR: File not found: {ui_spec_file}")
            print(f"   Current directory: {os.getcwd()}")
        sys.exit(1)
    
    create_ui_spec_docx(ui_spec_file, output_file)

if __name__ == '__main__':
    main()
