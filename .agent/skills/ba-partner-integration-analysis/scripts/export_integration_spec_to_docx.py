#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Script xuất Integration Specification ra file Word (.docx)
Hỗ trợ format tài liệu tích hợp đối tác chuẩn cho BA
"""

import sys
import os
import re
from pathlib import Path
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("Đang cài đặt thư viện python-docx...")
    os.system("pip install python-docx")
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

def set_cell_width(cell, width_inches):
    """Set width cho cell trong table"""
    cell_width = cell._element
    cell_width.set(qn('w:w'), str(int(width_inches * 1440)))
    cell_width.set(qn('w:type'), 'dxa')

def parse_markdown_table(content):
    """Parse bảng markdown format"""
    lines = content.strip().split('\n')
    headers = None
    rows = []
    in_table = False
    
    for line in lines:
        stripped = line.strip()
        
        # Bỏ qua separator line
        if stripped.startswith('|---') or re.match(r'^\|[-:|\s]+\|', stripped):
            continue
        
        # Parse header
        if '|' in line and headers is None:
            headers = [cell.strip() for cell in line.split('|') if cell.strip()]
            in_table = True
            continue
        
        # Parse rows
        if in_table and '|' in line:
            cells = [cell.strip() for cell in line.split('|') if cell.strip()]
            if cells and len(cells) == len(headers):
                rows.append(cells)
        elif in_table and not stripped:
            in_table = False
    
    if headers and rows:
        return {'headers': headers, 'rows': rows}
    return None

def extract_mermaid_code(content, diagram_type='sequence'):
    """Trích xuất Mermaid code từ content"""
    pattern = rf'```mermaid\s*\n(.*?)\n```'
    matches = re.findall(pattern, content, re.DOTALL | re.IGNORECASE)
    
    if matches:
        # Tìm diagram phù hợp với type
        for match in matches:
            if diagram_type == 'sequence' and 'sequenceDiagram' in match:
                return match.strip()
            elif diagram_type == 'flowchart' and ('flowchart' in match or 'graph' in match):
                return match.strip()
        
        # Nếu không tìm thấy type cụ thể, trả về match đầu tiên
        return matches[0].strip()
    
    return None

def extract_section(content, section_name):
    """Trích xuất nội dung của một section từ markdown"""
    pattern = rf'^##+\s*{re.escape(section_name)}'
    lines = content.split('\n')
    start_idx = None
    end_idx = None
    
    for i, line in enumerate(lines):
        if re.match(pattern, line, re.IGNORECASE):
            start_idx = i + 1
            break
    
    if start_idx is None:
        return None
    
    # Tìm section tiếp theo
    for i in range(start_idx, len(lines)):
        if re.match(r'^##+\s+', lines[i]):
            end_idx = i
            break
    
    if end_idx is None:
        end_idx = len(lines)
    
    return '\n'.join(lines[start_idx:end_idx]).strip()

def create_integration_spec_docx(integration_spec_file, output_file=None):
    """Tạo file Word từ Integration Spec"""
    
    # Đọc file
    with open(integration_spec_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Tạo Word document
    doc = Document()
    
    # Cấu hình font mặc định
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    # Title
    title = doc.add_heading('ĐẶC TẢ TÍCH HỢP ĐỐI TÁC (PARTNER INTEGRATION SPECIFICATION)', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Parse thông tin tích hợp
    integration_info = {}
    for line in content.split('\n'):
        line_stripped = line.strip()
        if re.match(r'^[Tt]ên tích hợp\s*[:：]', line_stripped):
            integration_info['name'] = re.split(r'[:：]', line_stripped, 1)[1].strip()
        elif re.match(r'^Mô tả\s*[:：]', line_stripped) or re.match(r'^Mục đích\s*[:：]', line_stripped):
            integration_info['description'] = re.split(r'[:：]', line_stripped, 1)[1].strip()
    
    # 1. Tên tích hợp, mô tả về mục đích tích hợp
    doc.add_paragraph()
    heading1 = doc.add_heading('1. Tên tích hợp', 1)
    
    if integration_info.get('name'):
        para = doc.add_paragraph()
        para.add_run(integration_info['name']).bold = True
    
    if integration_info.get('description'):
        doc.add_paragraph()
        para = doc.add_paragraph()
        para.add_run('Mô tả: ').bold = True
        para.add_run(integration_info['description'])
    
    # 2. Thông tin hệ thống đối tác
    doc.add_paragraph()
    heading2 = doc.add_heading('2. Thông tin hệ thống đối tác', 1)
    
    partner_section = extract_section(content, 'Hệ thống đối tác') or extract_section(content, 'Thông tin hệ thống đối tác')
    if partner_section:
        # Loại bỏ markdown formatting
        partner_text = re.sub(r'\*\*(.+?)\*\*', r'\1', partner_section)
        partner_text = re.sub(r'\*(.+?)\*', r'\1', partner_text)
        doc.add_paragraph(partner_text)
    else:
        doc.add_paragraph('(Chưa có thông tin hệ thống đối tác)')
    
    # 3. Nhu cầu tích hợp và dữ liệu đối tác cần
    doc.add_paragraph()
    heading3 = doc.add_heading('3. Nhu cầu tích hợp và dữ liệu đối tác cần', 1)
    
    requirement_section = extract_section(content, 'Nhu cầu tích hợp') or extract_section(content, 'Dữ liệu đối tác cần')
    if requirement_section:
        requirement_text = re.sub(r'\*\*(.+?)\*\*', r'\1', requirement_section)
        requirement_text = re.sub(r'\*(.+?)\*', r'\1', requirement_text)
        doc.add_paragraph(requirement_text)
    else:
        doc.add_paragraph('(Chưa có thông tin nhu cầu tích hợp)')
    
    # 4. Bảng phân tích dữ liệu đã có/chưa có
    doc.add_paragraph()
    heading4 = doc.add_heading('4. Bảng phân tích dữ liệu đã có/chưa có', 1)
    
    # Tìm bảng phân tích dữ liệu
    data_analysis_pattern = r'Tên trường thông tin.*?Mapping.*?Đã có'
    data_analysis_match = re.search(data_analysis_pattern, content, re.DOTALL | re.IGNORECASE)
    
    if data_analysis_match:
        # Tìm bảng tiếp theo
        table_start = data_analysis_match.end()
        table_end = content.find('\n\n', table_start)
        if table_end == -1:
            table_end = len(content)
        
        table_content = content[table_start:table_end]
        data_table = parse_markdown_table(table_content)
        if data_table:
            _add_table_to_doc(doc, data_table, 'Phân tích dữ liệu')
        else:
            doc.add_paragraph('(Chưa có bảng phân tích dữ liệu)')
    else:
        # Thử tìm theo section
        data_section = extract_section(content, 'Phân tích dữ liệu')
        if data_section:
            data_table = parse_markdown_table(data_section)
            if data_table:
                _add_table_to_doc(doc, data_table, 'Phân tích dữ liệu')
            else:
                doc.add_paragraph('(Chưa có bảng phân tích dữ liệu)')
        else:
            doc.add_paragraph('(Chưa có bảng phân tích dữ liệu)')
    
    # 5. Bảng phân tích chức năng cần thay đổi
    doc.add_paragraph()
    heading5 = doc.add_heading('5. Bảng phân tích chức năng cần thay đổi', 1)
    
    # Tìm bảng phân tích chức năng
    function_analysis_pattern = r'Tên chức năng.*?Loại thay đổi.*?Nội dung thay đổi'
    function_analysis_match = re.search(function_analysis_pattern, content, re.DOTALL | re.IGNORECASE)
    
    if function_analysis_match:
        # Tìm bảng tiếp theo
        table_start = function_analysis_match.end()
        table_end = content.find('\n\n', table_start)
        if table_end == -1:
            table_end = len(content)
        
        table_content = content[table_start:table_end]
        function_table = parse_markdown_table(table_content)
        if function_table:
            _add_table_to_doc(doc, function_table, 'Phân tích chức năng')
        else:
            doc.add_paragraph('(Chưa có bảng phân tích chức năng)')
    else:
        # Thử tìm theo section
        function_section = extract_section(content, 'Phân tích chức năng')
        if function_section:
            function_table = parse_markdown_table(function_section)
            if function_table:
                _add_table_to_doc(doc, function_table, 'Phân tích chức năng')
            else:
                doc.add_paragraph('(Chưa có bảng phân tích chức năng)')
        else:
            doc.add_paragraph('(Chưa có bảng phân tích chức năng)')
    
    # 6. Luồng tích hợp (Sequence diagram)
    doc.add_paragraph()
    heading6 = doc.add_heading('6. Luồng tích hợp (Sequence Diagram)', 1)
    
    sequence_code = extract_mermaid_code(content, 'sequence')
    if sequence_code:
        para = doc.add_paragraph()
        para.add_run('Mermaid Code:').bold = True
        doc.add_paragraph()
        code_para = doc.add_paragraph(sequence_code)
        code_para.style = 'No Spacing'
        # Format code với font monospace
        for run in code_para.runs:
            run.font.name = 'Courier New'
            run.font.size = Pt(10)
    else:
        doc.add_paragraph('(Chưa có sequence diagram)')
    
    # 7. Đặc tả API
    doc.add_paragraph()
    heading7 = doc.add_heading('7. Đặc tả API', 1)
    
    # Tìm tất cả các section API
    api_sections = re.findall(r'^##+\s*API\s*[:：]?\s*(.+?)$', content, re.MULTILINE | re.IGNORECASE)
    
    if not api_sections:
        # Thử tìm theo pattern khác
        api_sections = re.findall(r'^##+\s*(.+?API.+?)$', content, re.MULTILINE | re.IGNORECASE)
    
    # Parse toàn bộ content để tìm các API
    api_pattern = r'(?:^##+\s*API\s*[:：]?\s*(.+?)$|^###+\s*(.+?API.+?)$)(.*?)(?=^##+|\Z)'
    api_matches = re.findall(api_pattern, content, re.MULTILINE | re.DOTALL | re.IGNORECASE)
    
    if api_matches:
        for idx, match in enumerate(api_matches):
            api_name = match[0] or match[1] or f'API {idx+1}'
            api_content = match[2]
            
            # Sub-heading cho API
            doc.add_paragraph()
            api_heading = doc.add_heading(f'7.{idx+1} {api_name}', 2)
            
            # Tìm thông tin cơ bản
            basic_info_patterns = {
                'Mục đích': r'Mục đích\s*API[:\s]*\n(.*?)(?=\n\n|\nMethod|\Z)',
                'Method': r'Method[:\s]*\n(.*?)(?=\n\n|\nEndpoint|\Z)',
                'Endpoint': r'Endpoint[:\s]*\n(.*?)(?=\n\n|\nRequest|\Z)',
            }
            
            for info_name, pattern in basic_info_patterns.items():
                info_match = re.search(pattern, api_content, re.DOTALL | re.IGNORECASE)
                if info_match:
                    doc.add_paragraph()
                    para = doc.add_paragraph()
                    para.add_run(f'{info_name}: ').bold = True
                    para.add_run(info_match.group(1).strip())
            
            # Tìm Request Data
            request_match = re.search(r'Request Data[:\s]*\n(.*?)(?=\n\n|\nResponse|\nError|\nStatus|\Z)', api_content, re.DOTALL | re.IGNORECASE)
            if request_match:
                request_table = parse_markdown_table(request_match.group(1))
                if request_table:
                    _add_table_to_doc(doc, request_table, 'Request Data')
            
            # Tìm Response Data
            response_match = re.search(r'Response Data[:\s]*\n(.*?)(?=\n\n|\nError|\nStatus|\Z)', api_content, re.DOTALL | re.IGNORECASE)
            if response_match:
                response_table = parse_markdown_table(response_match.group(1))
                if response_table:
                    _add_table_to_doc(doc, response_table, 'Response Data')
            
            # Tìm Error Cases
            error_match = re.search(r'Error Case[:\s]*\n(.*?)(?=\n\n|\nStatus|\Z)', api_content, re.DOTALL | re.IGNORECASE)
            if error_match:
                doc.add_paragraph()
                error_heading = doc.add_heading('Các trường hợp lỗi', 3)
                error_text = error_match.group(1).strip()
                # Loại bỏ markdown formatting
                error_text = re.sub(r'\*\*(.+?)\*\*', r'\1', error_text)
                error_text = re.sub(r'\*(.+?)\*', r'\1', error_text)
                doc.add_paragraph(error_text)
            
            # Tìm Status Code
            status_match = re.search(r'Status Code[:\s]*\n(.*?)(?=\n\n|\Z)', api_content, re.DOTALL | re.IGNORECASE)
            if status_match:
                status_table = parse_markdown_table(status_match.group(1))
                if status_table:
                    _add_table_to_doc(doc, status_table, 'Status Code')
            
            # Tìm Security Considerations
            security_match = re.search(r'Security[:\s]*\n(.*?)(?=\n\n|\Z)', api_content, re.DOTALL | re.IGNORECASE)
            if security_match:
                doc.add_paragraph()
                security_heading = doc.add_heading('Security Considerations', 3)
                security_text = security_match.group(1).strip()
                security_text = re.sub(r'\*\*(.+?)\*\*', r'\1', security_text)
                security_text = re.sub(r'\*(.+?)\*', r'\1', security_text)
                doc.add_paragraph(security_text)
    
    # Lưu file
    if output_file is None:
        input_path = Path(integration_spec_file)
        output_file = str(input_path.with_suffix('.docx'))
    
    doc.save(output_file)
    try:
        print(f"SUCCESS: Đã tạo file Word: {output_file}")
    except UnicodeEncodeError:
        print(f"SUCCESS: Created Word file: {output_file}")
    return output_file

def _add_table_to_doc(doc, table_data, title=None):
    """Thêm bảng vào document"""
    if title:
        doc.add_paragraph()
        para = doc.add_paragraph()
        para.add_run(title).bold = True
    
    headers = table_data['headers']
    rows = table_data['rows']
    
    num_cols = len(headers)
    num_rows = len(rows) + 1  # +1 for header
    
    table = doc.add_table(rows=num_rows, cols=num_cols)
    table.style = 'Light Grid Accent 1'
    
    # Header
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_text = header.replace('**', '').replace('*', '').strip()
        header_cells[i].text = header_text
        
        # Bold header
        for paragraph in header_cells[i].paragraphs:
            for run in paragraph.runs:
                run.bold = True
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Rows
    for row_idx, row_data in enumerate(rows):
        row_cells = table.rows[row_idx + 1].cells
        for col_idx in range(num_cols):
            if col_idx < len(row_data):
                cell_text = row_data[col_idx]
                cell_text = cell_text.replace('**', '').replace('*', '').strip()
                row_cells[col_idx].text = cell_text
            else:
                row_cells[col_idx].text = ''

def main():
    if len(sys.argv) < 2:
        print("Usage: python export_integration_spec_to_docx.py <integration_spec_file> [output_file]")
        print("Example: python export_integration_spec_to_docx.py integration_spec.md")
        sys.exit(1)
    
    integration_spec_file_arg = sys.argv[1]
    
    # Xử lý đường dẫn
    integration_spec_file_path = Path(integration_spec_file_arg)
    if not integration_spec_file_path.is_absolute():
        integration_spec_file = str(integration_spec_file_path.resolve())
    else:
        integration_spec_file = str(integration_spec_file_path)
    
    if len(sys.argv) >= 3:
        output_file_arg = sys.argv[2]
        output_file_path = Path(output_file_arg)
        if not output_file_path.is_absolute():
            output_file = str(output_file_path.resolve())
        else:
            output_file = str(output_file_path)
    else:
        output_file = None
    
    if not os.path.exists(integration_spec_file):
        try:
            print(f"ERROR: Không tìm thấy file: {integration_spec_file}")
            print(f"   Đường dẫn hiện tại: {os.getcwd()}")
        except UnicodeEncodeError:
            print(f"ERROR: File not found: {integration_spec_file}")
            print(f"   Current directory: {os.getcwd()}")
        sys.exit(1)
    
    create_integration_spec_docx(integration_spec_file, output_file)

if __name__ == '__main__':
    main()
