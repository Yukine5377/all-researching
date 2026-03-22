#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Script xuất User Story/Use Case Documentation ra file Word (.docx)
Hỗ trợ format tài liệu User Story chuẩn cho BA với 7 mục
"""

import sys
import os
import re
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("Đang cài đặt thư viện python-docx...")
    os.system("pip install python-docx")
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH


def parse_markdown_table(content: str):
    """Parse bảng markdown format."""
    lines = [l for l in content.strip().split("\n") if l.strip()]
    headers = None
    rows = []
    in_table = False

    for line in lines:
        stripped = line.strip()

        # Bỏ qua separator line
        if stripped.startswith("|---") or re.match(r"^\|[-:|\s]+\|", stripped):
            continue

        if "|" in line and headers is None:
            headers = [c.strip() for c in line.split("|") if c.strip()]
            in_table = True
            continue

        if in_table and "|" in line:
            cells = [c.strip() for c in line.split("|") if c.strip()]
            if cells and len(cells) == len(headers):
                rows.append(cells)
        elif in_table and not stripped:
            break

    if headers and rows:
        return {"headers": headers, "rows": rows}
    return None


def extract_section_content(content: str, section_title: str):
    """Trích xuất nội dung của một section từ markdown."""
    # Tìm section (## hoặc ###)
    pattern = rf"^##+\s*{re.escape(section_title)}.*?\n(.*?)(?=^##|\Z)"
    match = re.search(pattern, content, re.MULTILINE | re.DOTALL | re.IGNORECASE)
    
    if match:
        return match.group(1).strip()
    return None


def extract_mermaid_diagram(content: str, diagram_type: str = None):
    """Trích xuất Mermaid diagram từ nội dung."""
    pattern = r"```mermaid\s*\n(.*?)\n```"
    matches = re.findall(pattern, content, re.DOTALL | re.IGNORECASE)
    
    if not matches:
        return None
    
    # Nếu có diagram_type, tìm diagram phù hợp
    if diagram_type:
        for m in matches:
            if diagram_type.lower() in m.lower():
                return m.strip()
    
    # Trả về diagram đầu tiên
    return matches[0].strip()


def extract_general_info(content: str):
    """Lấy thông tin mô tả chung (Tên chức năng, Actor, Mục đích)."""
    info = {}
    
    # Tìm "Tên chức năng:"
    match = re.search(r"Tên chức năng[:\s]+(.+?)(?:\n|$)", content, re.IGNORECASE)
    if match:
        info["name"] = match.group(1).strip()
    
    # Tìm "Actor:"
    match = re.search(r"Actor[:\s]+(.+?)(?:\n|$)", content, re.IGNORECASE)
    if match:
        info["actor"] = match.group(1).strip()
    
    # Tìm "Mục đích:"
    match = re.search(r"Mục đích[:\s]+(.+?)(?:\n|$)", content, re.IGNORECASE)
    if match:
        info["purpose"] = match.group(1).strip()
    
    return info


def extract_ui_spec_table(content: str):
    """Trích xuất bảng mô tả màn hình (UI Spec) từ section 3."""
    section_content = extract_section_content(content, "3. Mô tả màn hình")
    if not section_content:
        section_content = extract_section_content(content, "Mô tả màn hình")
    
    if section_content:
        # Tìm bảng đầu tiên trong section
        lines = section_content.split("\n")
        table_lines = []
        in_table = False
        
        for line in lines:
            if "|" in line:
                in_table = True
                table_lines.append(line)
            elif in_table and not line.strip():
                break
            elif in_table:
                table_lines.append(line)
        
        if table_lines:
            table_content = "\n".join(table_lines)
            return parse_markdown_table(table_content)
    
    return None


def extract_api_sections(content: str):
    """Trích xuất các section API từ markdown."""
    api_sections = []
    
    # Tìm phần "Danh sách các API" hoặc "Đặc tả API"
    api_section_start = re.search(r"^##+\s*(Danh sách các API|Đặc tả API)", content, re.MULTILINE | re.IGNORECASE)
    if not api_section_start:
        return api_sections
    
    # Tìm tất cả các API (### API: ... hoặc ### API 1: ...)
    api_pattern = r"^###+\s*API\s*(?:\d+)?\s*[:：]\s*(.+?)$"
    lines = content.split("\n")
    start_idx = api_section_start.start()
    
    current_api = None
    current_content = []
    
    for i in range(start_idx, len(lines)):
        line = lines[i]
        
        # Tìm API mới
        api_match = re.match(api_pattern, line, re.IGNORECASE)
        if api_match:
            # Lưu API cũ nếu có
            if current_api:
                api_sections.append({
                    "name": current_api,
                    "content": "\n".join(current_content)
                })
            
            # Bắt đầu API mới
            current_api = api_match.group(1).strip()
            current_content = []
        elif current_api:
            # Kiểm tra xem có phải section mới không (##)
            if re.match(r"^##+\s+", line):
                # Lưu API cuối cùng
                if current_api:
                    api_sections.append({
                        "name": current_api,
                        "content": "\n".join(current_content)
                    })
                break
            else:
                current_content.append(line)
    
    # Lưu API cuối cùng
    if current_api:
        api_sections.append({
            "name": current_api,
            "content": "\n".join(current_content)
        })
    
    return api_sections


def extract_api_basic_info(api_content: str):
    """Trích xuất thông tin cơ bản của API."""
    info = {}
    
    # Tìm Tên API
    name_match = re.search(r"\*\*Tên API:\*\*\s*(.+?)(?:\n|$)", api_content, re.IGNORECASE)
    if name_match:
        info["name"] = name_match.group(1).strip()
    
    # Tìm Mục đích
    purpose_match = re.search(r"\*\*Mục đích:\*\*\s*(.+?)(?:\n|$)", api_content, re.IGNORECASE)
    if purpose_match:
        info["purpose"] = purpose_match.group(1).strip()
    
    # Tìm Method
    method_match = re.search(r"\*\*Method:\*\*\s*(.+?)(?:\n|$)", api_content, re.IGNORECASE)
    if method_match:
        info["method"] = method_match.group(1).strip()
    
    # Tìm Endpoint
    endpoint_match = re.search(r"\*\*Endpoint:\*\*\s*(.+?)(?:\n|$)", api_content, re.IGNORECASE)
    if endpoint_match:
        info["endpoint"] = endpoint_match.group(1).strip()
    
    return info


def extract_table_by_title(content: str, title: str):
    """Trích xuất bảng theo title (Request Data, Response Data, Status Code)."""
    pattern = rf"\*\*{re.escape(title)}\*\*[:\s]*\n(.*?)(?=\n\*\*|\n##|\Z)"
    match = re.search(pattern, content, re.DOTALL | re.IGNORECASE)
    
    if match:
        table_content = match.group(1)
        return parse_markdown_table(table_content)
    return None


def extract_error_cases(api_content: str):
    """Trích xuất các trường hợp lỗi."""
    pattern = r"\*\*Các trường hợp lỗi.*?\*\*[:\s]*\n(.*?)(?=\n\*\*Status Code|\n\*\*|\Z)"
    match = re.search(pattern, api_content, re.DOTALL | re.IGNORECASE)
    
    if match:
        error_text = match.group(1).strip()
        error_text = re.sub(r"\*\*(.+?)\*\*", r"\1", error_text)
        error_text = re.sub(r"\*(.+?)\*", r"\1", error_text)
        return error_text
    return None


def extract_activity_flow_table(content: str):
    """Trích xuất bảng mô tả luồng theo Activity Diagram từ section 7."""
    section_content = extract_section_content(content, "7. Bảng mô tả luồng")
    if not section_content:
        section_content = extract_section_content(content, "Bảng mô tả luồng")
    
    if section_content:
        # Tìm bảng đầu tiên trong section
        return parse_markdown_table(section_content)
    
    return None


def add_table(doc: Document, table_data: dict, title: str = None):
    """Thêm bảng vào Word."""
    if title:
        p = doc.add_paragraph()
        p.add_run(title).bold = True

    headers = table_data["headers"]
    rows = table_data["rows"]

    tbl = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    tbl.style = "Light Grid Accent 1"

    # Header
    hdr_cells = tbl.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for para in hdr_cells[i].paragraphs:
            for run in para.runs:
                run.bold = True
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Rows
    for r_idx, row in enumerate(rows):
        row_cells = tbl.rows[r_idx + 1].cells
        for c_idx, value in enumerate(row):
            row_cells[c_idx].text = value


def create_user_story_docx(input_file: str, output_file: str = None):
    """Tạo file Word từ file markdown User Story."""
    with open(input_file, "r", encoding="utf-8") as f:
        content = f.read()

    doc = Document()

    # Font mặc định
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    # Title
    title = doc.add_heading("USER STORY / USE CASE DOCUMENTATION", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 1. Mô tả chung
    doc.add_paragraph()
    doc.add_heading("1. Mô tả chung", level=1)
    
    general_info = extract_general_info(content)
    
    if general_info.get("name"):
        p = doc.add_paragraph()
        p.add_run("Tên chức năng: ").bold = True
        p.add_run(general_info["name"])
    
    if general_info.get("actor"):
        p = doc.add_paragraph()
        p.add_run("Actor: ").bold = True
        p.add_run(general_info["actor"])
    
    if general_info.get("purpose"):
        p = doc.add_paragraph()
        p.add_run("Mục đích: ").bold = True
        p.add_run(general_info["purpose"])

    # 2. Màn hình
    doc.add_paragraph()
    doc.add_heading("2. Màn hình", level=1)
    doc.add_paragraph("(Người dùng tự chèn ảnh màn hình vào đây)")

    # 3. Mô tả màn hình
    doc.add_paragraph()
    doc.add_heading("3. Mô tả màn hình", level=1)
    
    ui_spec_table = extract_ui_spec_table(content)
    if ui_spec_table:
        add_table(doc, ui_spec_table, title="Bảng đặc tả màn hình")
    else:
        doc.add_paragraph("(Chưa có bảng mô tả màn hình)")

    # 4. Sequence Diagram
    doc.add_paragraph()
    doc.add_heading("4. Sequence Diagram", level=1)
    
    seq_code = extract_mermaid_diagram(content, "sequenceDiagram")
    if seq_code:
        p = doc.add_paragraph()
        p.add_run("Mermaid Code:").bold = True
        code_para = doc.add_paragraph()
        code_para.style = "No Spacing"
        run = code_para.add_run(seq_code)
        run.font.name = "Courier New"
        run.font.size = Pt(10)
    else:
        doc.add_paragraph("(Chưa có Sequence Diagram trong file)")

    # 5. Danh sách các API
    doc.add_paragraph()
    doc.add_heading("5. Danh sách các API", level=1)
    
    api_sections = extract_api_sections(content)
    if api_sections:
        for idx, api_section in enumerate(api_sections, 1):
            api_name = api_section["name"]
            api_content = api_section["content"]
            
            # Sub-heading cho API
            doc.add_paragraph()
            doc.add_heading(f"5.{idx} API: {api_name}", level=2)
            
            # Thông tin cơ bản
            basic_info = extract_api_basic_info(api_content)
            if basic_info:
                doc.add_paragraph()
                p = doc.add_paragraph()
                p.add_run("Thông tin cơ bản:").bold = True
                
                if basic_info.get("name"):
                    doc.add_paragraph(f"Tên API: {basic_info['name']}")
                if basic_info.get("purpose"):
                    doc.add_paragraph(f"Mục đích: {basic_info['purpose']}")
                if basic_info.get("method"):
                    doc.add_paragraph(f"Method: {basic_info['method']}")
                if basic_info.get("endpoint"):
                    doc.add_paragraph(f"Endpoint: {basic_info['endpoint']}")
            
            # Request Data
            request_table = extract_table_by_title(api_content, "Request Data")
            if request_table:
                doc.add_paragraph()
                add_table(doc, request_table, title="Request Data")
            
            # Response Data
            response_table = extract_table_by_title(api_content, "Response Data")
            if response_table:
                doc.add_paragraph()
                add_table(doc, response_table, title="Response Data")
            
            # Error Cases
            error_cases = extract_error_cases(api_content)
            if error_cases:
                doc.add_paragraph()
                p = doc.add_paragraph()
                p.add_run("Các trường hợp lỗi:").bold = True
                doc.add_paragraph(error_cases)
            
            # Status Code
            status_table = extract_table_by_title(api_content, "Status Code")
            if status_table:
                doc.add_paragraph()
                add_table(doc, status_table, title="Status Code")
    else:
        doc.add_paragraph("(Chưa có đặc tả API trong file)")

    # 6. Activity Diagram
    doc.add_paragraph()
    doc.add_heading("6. Activity Diagram", level=1)
    doc.add_paragraph("(Người dùng tự mở file .drawio trong Draw.io để xem Activity Diagram)")

    # 7. Bảng mô tả luồng theo Activity Diagram
    doc.add_paragraph()
    doc.add_heading("7. Bảng mô tả luồng theo Activity Diagram", level=1)
    
    activity_flow_table = extract_activity_flow_table(content)
    if activity_flow_table:
        add_table(doc, activity_flow_table, title="Bảng mô tả luồng")
    else:
        doc.add_paragraph("(Chưa có bảng mô tả luồng trong file)")

    # Lưu file
    if output_file is None:
        path = Path(input_file)
        output_file = str(path.with_suffix(".docx"))

    doc.save(output_file)
    try:
        print(f"SUCCESS: Đã tạo file Word: {output_file}")
    except UnicodeEncodeError:
        print(f"SUCCESS: Created Word file: {output_file}")

    return output_file


def main():
    if len(sys.argv) < 2:
        print("Usage: python export_user_story_to_docx.py <user_story_file> [output_file]")
        sys.exit(1)

    input_arg = sys.argv[1]
    input_path = Path(input_arg)
    if not input_path.is_absolute():
        input_file = str(input_path.resolve())
    else:
        input_file = str(input_path)

    if len(sys.argv) >= 3:
        out_arg = sys.argv[2]
        out_path = Path(out_arg)
        if not out_path.is_absolute():
            output_file = str(out_path.resolve())
        else:
            output_file = str(out_path)
    else:
        output_file = None

    if not os.path.exists(input_file):
        print(f"ERROR: Không tìm thấy file: {input_file}")
        print(f"   Đường dẫn hiện tại: {os.getcwd()}")
        sys.exit(1)

    create_user_story_docx(input_file, output_file)


if __name__ == "__main__":
    main()
