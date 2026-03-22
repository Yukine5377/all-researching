#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Script xuất Integration Specification ra file Word (.docx)
Hỗ trợ format tài liệu tích hợp API chuẩn cho BA
"""

import sys
import os
import re
from pathlib import Path
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("Đang cài đặt thư viện python-docx...")
    os.system("pip install python-docx")
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH


def parse_markdown_table(content: str):
    """Parse bảng markdown format đơn giản."""
    lines = [l for l in content.strip().split("\n") if l.strip()]
    headers = None
    rows = []
    in_table = False

    for line in lines:
        stripped = line.strip()

        # Bỏ qua separator line
        if stripped.startswith("|---") or re.match(r"^\|[-:|\s]+\|", stripped):
            continue

        if "|" in line and headers is None:
            headers = [c.strip() for c in line.split("|") if c.strip()]
            in_table = True
            continue

        if in_table and "|" in line:
            cells = [c.strip() for c in line.split("|") if c.strip()]
            if cells and len(cells) == len(headers):
                rows.append(cells)
        elif in_table and not stripped:
            break

    if headers and rows:
        return {"headers": headers, "rows": rows}
    return None


def extract_mermaid_diagram(content: str, diagram_type: str = "sequenceDiagram"):
    """Trích xuất Mermaid diagram (sequenceDiagram hoặc erDiagram)."""
    pattern = rf"```mermaid\s*\n(.*?)\n```"
    matches = re.findall(pattern, content, re.DOTALL | re.IGNORECASE)
    if not matches:
        return None
    for m in matches:
        if diagram_type.lower() in m.lower():
            return m.strip()
    return None


def extract_section(content: str, section_title: str):
    """Trích xuất nội dung section theo title (A, B, C, D, E, F, G, H)."""
    # Pattern để tìm section: ## A. hoặc ## B. hoặc ### A. ...
    pattern = rf"^##+\s*{re.escape(section_title)}\.\s*(.+?)$"
    lines = content.split("\n")
    start_idx = None
    
    for i, line in enumerate(lines):
        if re.match(pattern, line, re.IGNORECASE):
            start_idx = i
            break
    
    if start_idx is None:
        return None
    
    # Tìm section tiếp theo hoặc hết file
    end_idx = len(lines)
    for i in range(start_idx + 1, len(lines)):
        if re.match(r"^##+\s*[A-H]\.\s*", lines[i], re.IGNORECASE):
            end_idx = i
            break
    
    return "\n".join(lines[start_idx:end_idx])


def extract_system_name(content: str):
    """Lấy tên hệ thống từ đầu file hoặc từ path."""
    # Tìm trong content
    match = re.search(r"Tên hệ thống[:\s]+(.+?)(?:\n|$)", content, re.IGNORECASE)
    if match:
        return match.group(1).strip()
    
    # Hoặc tìm trong path
    return None


def add_table(doc: Document, table_data: dict, title: str = None):
    """Thêm bảng vào Word."""
    if not table_data or not table_data.get("headers"):
        return
    
    if title:
        p = doc.add_paragraph()
        p.add_run(title).bold = True

    headers = table_data["headers"]
    rows = table_data.get("rows", [])

    tbl = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    tbl.style = "Light Grid Accent 1"

    # Header
    hdr_cells = tbl.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for para in hdr_cells[i].paragraphs:
            for run in para.runs:
                run.bold = True
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Rows
    for r_idx, row in enumerate(rows):
        if r_idx + 1 >= len(tbl.rows):
            break
        row_cells = tbl.rows[r_idx + 1].cells
        for c_idx, value in enumerate(row):
            if c_idx < len(row_cells):
                row_cells[c_idx].text = value


def add_markdown_content(doc: Document, content: str):
    """Thêm nội dung markdown vào Word (đơn giản, không parse hết)."""
    lines = content.split("\n")
    i = 0
    
    while i < len(lines):
        line = lines[i].strip()
        
        # Skip empty
        if not line:
            i += 1
            continue
        
        # Heading
        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
        # Table
        elif "|" in line and i + 1 < len(lines) and "---" in lines[i + 1]:
            # Parse table
            table_lines = [line]
            i += 1
            while i < len(lines) and ("|" in lines[i] or not lines[i].strip()):
                if lines[i].strip():
                    table_lines.append(lines[i])
                i += 1
            table_content = "\n".join(table_lines)
            table_data = parse_markdown_table(table_content)
            if table_data:
                add_table(doc, table_data)
            continue
        # Paragraph
        else:
            # Remove markdown formatting đơn giản
            text = re.sub(r"\*\*(.+?)\*\*", r"\1", line)
            text = re.sub(r"\*(.+?)\*", r"\1", text)
            text = re.sub(r"`(.+?)`", r"\1", text)
            if text.strip():
                doc.add_paragraph(text)
        
        i += 1


def create_integration_spec_docx(input_file: str, output_file: str = None):
    """Tạo file Word từ file markdown integration spec."""
    with open(input_file, "r", encoding="utf-8") as f:
        content = f.read()

    doc = Document()

    # Font mặc định
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    # Title
    system_name = extract_system_name(content) or "Hệ thống đối tác"
    title = doc.add_heading(f"TÀI LIỆU TÍCH HỢP API - {system_name.upper()}", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Date
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run(f"Ngày tạo: {datetime.now().strftime('%d/%m/%Y')}").italic = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # A. Tổng quan & Mục đích tích hợp
    doc.add_page_break()
    doc.add_heading("A. Tổng quan & Mục đích tích hợp", level=1)
    section_a = extract_section(content, "A")
    if section_a:
        add_markdown_content(doc, section_a)
    else:
        doc.add_paragraph("(Chưa có nội dung)")

    # B. Luồng tích hợp tổng thể
    doc.add_paragraph()
    doc.add_heading("B. Luồng tích hợp tổng thể", level=1)
    section_b = extract_section(content, "B")
    if section_b:
        add_markdown_content(doc, section_b)
    else:
        doc.add_paragraph("(Chưa có nội dung)")

    # C. Sequence diagram
    doc.add_paragraph()
    doc.add_heading("C. Sequence diagram", level=1)
    seq_code = extract_mermaid_diagram(content, "sequenceDiagram")
    if seq_code:
        p = doc.add_paragraph()
        p.add_run("Mermaid Code:").bold = True
        code_para = doc.add_paragraph()
        code_para.style = "No Spacing"
        run = code_para.add_run(seq_code)
        run.font.name = "Courier New"
        run.font.size = Pt(9)
    else:
        doc.add_paragraph("(Chưa có Mermaid sequenceDiagram)")

    # D. Bảng phân tích API chi tiết
    doc.add_paragraph()
    doc.add_heading("D. Bảng phân tích API chi tiết", level=1)
    section_d = extract_section(content, "D")
    if section_d:
        add_markdown_content(doc, section_d)
    else:
        doc.add_paragraph("(Chưa có nội dung)")

    # D'. Bảng mapping dữ liệu request
    doc.add_paragraph()
    doc.add_heading("D'. Bảng mapping dữ liệu request", level=1)
    # Tìm section D' trong content
    d_prime_pattern = r"^##+\s*D'\.\s*(.+?)(?=^##+\s*[E-H]\.|$)"
    d_prime_match = re.search(d_prime_pattern, content, re.MULTILINE | re.DOTALL | re.IGNORECASE)
    if d_prime_match:
        add_markdown_content(doc, d_prime_match.group(0))
    else:
        doc.add_paragraph("(Chưa có nội dung)")

    # E. Bảng mapping dữ liệu response
    doc.add_paragraph()
    doc.add_heading("E. Bảng mapping dữ liệu response → model nội bộ", level=1)
    section_e = extract_section(content, "E")
    if section_e:
        add_markdown_content(doc, section_e)
    else:
        doc.add_paragraph("(Chưa có nội dung)")

    # F. Danh sách Use Case
    doc.add_paragraph()
    doc.add_heading("F. Danh sách Use Case & màn hình/chức năng cần xây", level=1)
    section_f = extract_section(content, "F")
    if section_f:
        add_markdown_content(doc, section_f)
    else:
        doc.add_paragraph("(Chưa có nội dung)")

    # G. ERD
    doc.add_paragraph()
    doc.add_heading("G. ERD (Mermaid erDiagram)", level=1)
    erd_code = extract_mermaid_diagram(content, "erDiagram")
    if erd_code:
        p = doc.add_paragraph()
        p.add_run("Mermaid Code:").bold = True
        code_para = doc.add_paragraph()
        code_para.style = "No Spacing"
        run = code_para.add_run(erd_code)
        run.font.name = "Courier New"
        run.font.size = Pt(9)
    else:
        doc.add_paragraph("(Chưa có Mermaid erDiagram)")

    # H. Checklist
    doc.add_paragraph()
    doc.add_heading("H. Checklist cho BA để hoàn thiện & go-live", level=1)
    section_h = extract_section(content, "H")
    if section_h:
        add_markdown_content(doc, section_h)
    else:
        doc.add_paragraph("(Chưa có nội dung)")

    if output_file is None:
        path = Path(input_file)
        output_file = str(path.with_suffix(".docx"))

    doc.save(output_file)
    try:
        print(f"SUCCESS: Đã tạo file Word: {output_file}")
    except UnicodeEncodeError:
        print(f"SUCCESS: Created Word file: {output_file}")

    return output_file


def main():
    if len(sys.argv) < 2:
        print("Usage: python export_integration_spec_to_docx.py <integration_spec_file> [output_file]")
        sys.exit(1)

    input_arg = sys.argv[1]
    input_path = Path(input_arg)
    if not input_path.is_absolute():
        input_file = str(input_path.resolve())
    else:
        input_file = str(input_path)

    if len(sys.argv) >= 3:
        out_arg = sys.argv[2]
        out_path = Path(out_arg)
        if not out_path.is_absolute():
            output_file = str(out_path.resolve())
        else:
            output_file = str(out_path)
    else:
        output_file = None

    if not os.path.exists(input_file):
        print(f"ERROR: Không tìm thấy file: {input_file}")
        print(f"   Đường dẫn hiện tại: {os.getcwd()}")
        sys.exit(1)

    create_integration_spec_docx(input_file, output_file)


if __name__ == "__main__":
    main()
